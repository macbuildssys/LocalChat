"""
Extract plain text from uploaded files.
Each parser is imported lazily so missing optional deps don't crash the server.
"""
import os
import tempfile
from pathlib import Path

DOCUMENT_EXTENSIONS = {
    ".pdf", ".docx", ".doc", ".epub", ".odt", ".ods", ".odp",
    ".txt", ".md", ".rst", ".csv", ".json", ".xml",
    ".html", ".htm", ".rtf",
}

IMAGE_EXTENSIONS = {
    ".jpg", ".jpeg", ".png", ".gif", ".webp",
    ".bmp", ".tiff", ".tif", ".svg",
}

def is_image(filename: str) -> bool:
    return Path(filename).suffix.lower() in IMAGE_EXTENSIONS

def is_document(filename: str) -> bool:
    return Path(filename).suffix.lower() in DOCUMENT_EXTENSIONS

def extract_text(content: bytes, filename: str) -> str:
    suffix = Path(filename).suffix.lower()
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as f:
        f.write(content)
        tmp = f.name
    try:
        return _dispatch(tmp, suffix)
    finally:
        try:
            os.unlink(tmp)
        except OSError:
            pass

def _dispatch(path: str, suffix: str) -> str:
    if suffix == ".pdf":
        return _pdf(path)
    if suffix in (".docx", ".doc"):
        return _docx(path)
    if suffix == ".epub":
        return _epub(path)
    if suffix in (".odt", ".ods", ".odp"):
        return _odt(path)
    if suffix == ".rtf":
        return _rtf(path)
    if suffix in (".txt", ".md", ".rst", ".csv", ".json", ".xml", ".html", ".htm"):
        with open(path, "r", errors="replace") as fh:
            return fh.read()
    raise ValueError(f"Unsupported file type: {suffix}")

def _pdf(path: str) -> str:
    try:
        import fitz  # pymupdf
        doc = fitz.open(path)
        pages = [page.get_text() for page in doc]
        doc.close()
        return "\n\n".join(p for p in pages if p.strip())
    except ImportError:
        raise RuntimeError("pymupdf not installed. Run: pip install pymupdf")

def _docx(path: str) -> str:
    try:
        from docx import Document
        doc = Document(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                paragraphs.append("\t".join(cell.text for cell in row.cells))
        return "\n".join(paragraphs)
    except ImportError:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")

def _epub(path: str) -> str:
    try:
        import ebooklib
        from ebooklib import epub
        from bs4 import BeautifulSoup
    except ImportError:
        raise RuntimeError(
            "Missing deps. Run: pip install ebooklib beautifulsoup4"
        )
    book = epub.read_epub(path, options={"ignore_ncx": True})
    texts = []
    for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
        soup = BeautifulSoup(item.get_content(), "html.parser")
        t = soup.get_text(separator="\n", strip=True)
        if t:
            texts.append(t)
    return "\n\n".join(texts)

def _odt(path: str) -> str:
    try:
        from odf import teletype
        from odf.opendocument import load
    except ImportError:
        raise RuntimeError("odfpy not installed. Run: pip install odfpy")
    doc = load(path)
    return teletype.extractText(doc.text)

def _rtf(path: str) -> str:
    # Basic RTF stripping — good enough for plain text extraction
    with open(path, "rb") as f:
        raw = f.read().decode("latin-1", errors="replace")
    import re
    # Remove RTF control words and groups
    text = re.sub(r"\\\w+\*?", " ", raw)
    text = re.sub(r"[{}\\]", "", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()
